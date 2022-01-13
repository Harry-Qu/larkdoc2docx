import os.path
import re

from docx import Document
from docx.styles.style import _ParagraphStyle
from docx.enum.style import WD_STYLE_TYPE


class larkDoc2Docx(object):
    styles = []
    USED_STYLE_NAME = ['Title', 'Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
                       'Heading 7', 'Heading 8', 'Heading 9', 'List Paragraph', '图片', '代码', '表格正文']

    @staticmethod
    def print_doc_paragraph_style(document):
        """
        打印文档每段的样式

        :param document:文档对象
        """
        paragraphs = document.paragraphs
        for paragraph in paragraphs:
            print(paragraph.style)

    def read_template_style(self, filename: str):
        """
        读取文档样式列表

        :param filename:文件名
        """

        if not os.path.isfile(filename):
            print("未找到模版文档")
            return []
        document = Document(filename)
        styles = document.styles
        # s = []
        for style in styles:
            if (type(style) != _ParagraphStyle):
                continue
        #     s.append((style.name, style.type))
        # print(s)
        self.styles = styles
        return styles

    def add_styles_to_document(self, document):
        """
        为文档添加样式

        :param document: 文档对象
        :param document: 表格对象
        """

        styles = document.styles
        style_names = []

        for style in self.styles:
            if not (style.type in [WD_STYLE_TYPE.PARAGRAPH]):
                continue
            if styles.__contains__(style.name):
                continue
            new_style = styles.add_style(style.name, style.type)
            style_names.append(style.name)

            if style.type in [WD_STYLE_TYPE.PARAGRAPH]:
                # new_style.font.element=style.font.element
                new_style.font.bold = style.font.bold
                new_style.font.outline = style.font.outline
                new_style.font.size = style.font.size
                new_style.font.all_caps = style.font.all_caps
                new_style.font.complex_script = style.font.complex_script
                new_style.font.cs_bold = style.font.cs_bold
                new_style.font.cs_italic = style.font.cs_italic
                new_style.font.double_strike = style.font.double_strike
                new_style.font.emboss = style.font.emboss
                new_style.font.hidden = style.font.hidden
                new_style.font.highlight_color = style.font.highlight_color
                new_style.font.imprint = style.font.imprint
                new_style.font.italic = style.font.italic
                new_style.font.small_caps = style.font.small_caps
                new_style.font.snap_to_grid = style.font.snap_to_grid
                new_style.font.spec_vanish = style.font.spec_vanish
                new_style.font.name = style.font.name
                try:
                    new_style.font.chinese_name = style.font.chinese_name
                except AttributeError:
                    pass
                new_style.paragraph_format.alignment = style.paragraph_format.alignment
                new_style.paragraph_format.first_line_indent = style.paragraph_format.first_line_indent
                new_style.paragraph_format.keep_together = style.paragraph_format.keep_together
                new_style.paragraph_format.keep_with_next = style.paragraph_format.keep_with_next
                new_style.paragraph_format.left_indent = style.paragraph_format.left_indent
                new_style.paragraph_format.line_spacing = style.paragraph_format.line_spacing
                new_style.paragraph_format.line_spacing_rule = style.paragraph_format.line_spacing_rule
                new_style.paragraph_format.page_break_before = style.paragraph_format.page_break_before
                new_style.paragraph_format.right_indent = style.paragraph_format.right_indent
                new_style.paragraph_format.space_after = style.paragraph_format.space_after
                new_style.paragraph_format.space_before = style.paragraph_format.space_before
                new_style.paragraph_format.widow_control = style.paragraph_format.widow_control

        for used_style in self.USED_STYLE_NAME:
            if not styles.__contains__(used_style):
                styles.add_style(used_style, WD_STYLE_TYPE.PARAGRAPH)

        print("已添加样式:{}".format(",".join(style_names)))
        return document

    def _change_table_style(self, document, table):
        """
        更改表格的样式

        :param document: 文档对象
        :param document: 表格对象
        """
        code_res = re.search('w:w="0"', table._element.xml)  # 判断是否为代码
        rows = table.rows

        for row in rows:
            for cell in row.cells:
                cell_paragraphs = cell.paragraphs
                for cell_paragraph in cell_paragraphs:
                    if code_res is None:
                        print("为表格添加样式:表格正文")
                        cell_paragraph.style = document.styles["表格正文"]
                    else:
                        print("为表格添加样式:代码")
                        cell_paragraph.style = document.styles["代码"]
                    self._clear_paragraph_specific_style(cell_paragraph)

    @staticmethod
    def _get_style_by_outline(document, outline_level: int):
        """
        根据大纲级别获取样式

        :param outline_level: 大纲级别（飞书文档中标题1的级别为0）
        """
        return document.styles["Heading {}".format(outline_level)]

    def _get_paragraph_style(self, document, paragraph):
        """
        根据段落内容获取样式

        :param document: 文档对象
        :param paragraph: 段落对象
        """
        p = paragraph._p

        # print(paragraph.runs[0].font.outline)

        # 为标题1-标题9添加样式
        outline_res = re.search('<w:outlineLvl w:val="(\d+)"/>', p.xml)
        if outline_res is not None:
            outline_level = int(outline_res.group(1)) + 1
            print("为段落添加样式:标题{}".format(outline_level))
            return self._get_style_by_outline(document, outline_level)

        # 为标题添加样式
        if paragraph.paragraph_format.space_after is not None and \
                paragraph.paragraph_format.space_after.twips == 480:
            print("为段落添加样式:标题")
            return document.styles["Title"]

        # 为图片添加样式
        pic_res = re.search('<w:drawing>', p.xml)
        if pic_res is not None:
            print("为段落添加样式:图片")
            return document.styles["图片"]

        # 为列表段落添加样式
        list_res = re.search('<w:numPr>', p.xml)
        if list_res is not None:
            print("为段落添加样式:列表")
            return document.styles["List Paragraph"]

        # 其他类型默认为正文
        return document.styles["Normal"]

    @staticmethod
    def _clear_paragraph_specific_style(paragraph):
        """
        清除原有的格式

        :param paragraph: 要清除格式的段落
        """
        for run in paragraph.runs:
            run.style = None
            run.font.name = None
            run.font.size = None
            run.font.bold = None
            run.font.name = None
            try:
                run.font.chinese_name = None
            except AttributeError:
                pass
        paragraph.paragraph_format.alignment = None
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.keep_together = None
        paragraph.paragraph_format.keep_with_next = None
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.line_spacing = None
        paragraph.paragraph_format.line_spacing_rule = None
        paragraph.paragraph_format.page_break_before = None
        paragraph.paragraph_format.right_indent = None
        paragraph.paragraph_format.space_after = None
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.widow_control = None

    def change_larkdoc_style(self, document):
        """
        更改飞书文档各段落的样式

        :param document: 文档对象
        """

        # 为文字段落修改样式
        paragraphs = document.paragraphs
        for paragraph in paragraphs:
            paragraph.style = self._get_paragraph_style(document, paragraph)
            self._clear_paragraph_specific_style(paragraph)

        # 为表格段落修改样式
        for table in document.tables:
            self._change_table_style(document, table)

        return document
